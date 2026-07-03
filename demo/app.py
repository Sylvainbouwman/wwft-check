import streamlit as st
import streamlit.components.v1 as components
import requests
import os
import pandas as pd
from datetime import datetime
from io import BytesIO
from dotenv import load_dotenv
from duckduckgo_search import DDGS
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()

KVK_API_KEY = os.getenv("KVK_API_KEY", "")
KVK_BASE_URL = os.getenv("KVK_BASE_URL", "https://api.kvk.nl/api/v2")
OPENSANCTIONS_API_KEY = os.getenv("OPENSANCTIONS_API_KEY", "")
DEMO_MODUS = not bool(KVK_API_KEY)

_kvk_session = requests.Session()
_kvk_session.headers.update({
    "apikey": KVK_API_KEY,
    "Accept": "application/json",
    "Connection": "close",
})

st.set_page_config(
    page_title="WWFT Check Tool",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
/* Metrics */
[data-testid="stMetricValue"] { font-size: 1rem !important; }
[data-testid="stMetricLabel"] { font-size: 0.78rem !important; color: #6c757d; }

/* Alerts & spacing */
.stAlert { margin-top: 0.4rem; border-radius: 8px !important; }

/* Data editor */
[data-testid="stDataEditorContainer"] { border-radius: 8px; overflow: hidden; border: 1px solid #dee2e6; }

/* Forms */
[data-testid="stForm"] { border: none !important; padding: 0 !important; }

/* Download button full width */
[data-testid="stDownloadButton"] > button { width: 100%; font-size: 1rem; padding: 0.6rem 1.2rem; }

/* Divider spacing */
hr { margin: 1.5rem 0 !important; }

/* Step completed summary blocks */
.summary-block {
    background: #f8f9fa;
    border-left: 4px solid #1D9E75;
    border-radius: 0 8px 8px 0;
    padding: 0.6rem 1rem;
    margin-bottom: 0.5rem;
    font-size: 0.9rem;
}

/* Sidebar progress */
.prog-item { display: flex; align-items: center; gap: 8px; padding: 3px 0; font-size: 0.82rem; line-height: 1.4; }
.prog-done { color: #1D9E75; }
.prog-todo { color: #adb5bd; }
.prog-active { color: #378ADD; font-weight: 600; }
</style>
""", unsafe_allow_html=True)

# ──────────────────────────────────────────────
# Demo data
# ──────────────────────────────────────────────

DEMO_KVK_DATA = {
    "kvkNummer": "34220805",
    "naam": "Demo Vastgoed Amsterdam BV",
    "type": "hoofdvestiging",
    "adres": {
        "binnenlandsAdres": {
            "straatnaam": "Herengracht",
            "huisnummer": 182,
            "postcode": "1016BR",
            "plaats": "Amsterdam",
        }
    },
}

DEMO_BASISPROFIEL = {
    "kvkNummer": "34220805",
    "naam": "Demo Vastgoed Amsterdam BV",
    "sbiActiviteiten": [
        {"sbiCode": "6810", "sbiOmschrijving": "Handel in eigen onroerend goed"},
        {"sbiCode": "6820", "sbiOmschrijving": "Verhuur van en handel in eigen onroerend goed"},
    ],
    "totaalWerkzamePersonen": 7,
    "_embedded": {
        "eigenaar": {"rechtsvorm": "Besloten Vennootschap"},
        "hoofdvestiging": {
            "adressen": [{"volledigAdres": "Herengracht 182, 1016BR Amsterdam"}],
            "websites": ["www.demovastgoed.nl"],
        },
    },
    "handelsnamen": [{"naam": "Demo Vastgoed Amsterdam BV"}, {"naam": "DVA Vastgoed"}],
    "formeleRegistratiedatum": "20100315",
}


# ──────────────────────────────────────────────
# Constanten
# ──────────────────────────────────────────────

HOOG_RISICO_SBI = {"6810", "6820", "6492", "6491", "9200", "9201", "9202", "6612", "6619"}
MIDDEN_RISICO_SBI = {"5610", "5630", "6411", "6419", "6499", "7711", "7712"}
HOOG_RISICO_RECHTSVORM = {"Stichting", "Cooperatie", "Commanditaire Vennootschap"}

EU_HOOG_RISICO_LANDEN = {
    "Afghanistan", "Barbados", "Burkina Faso", "Cameroon", "Cayman Islands",
    "Congo (DRC)", "Gibraltar", "Haiti", "Jamaica", "Jordanie",
    "Mali", "Mozambique", "Myanmar", "Nicaragua", "Nigeria",
    "Pakistan", "Panama", "Filipijnen", "Senegal", "Zuid-Afrika",
    "Zuid-Soedan", "Syrie", "Tanzania", "Trinidad en Tobago", "Uganda",
    "Verenigde Arabische Emiraten", "Vanuatu", "Vietnam", "Jemen",
    "Noord-Korea", "Iran",
}

DIENSTVERLENING_OPTIES = [
    "Accountancy / jaarrekening",
    "Belastingadvies / aangifte",
    "Salarisadministratie",
    "Administratie / boekhouding",
    "Advies / consultancy",
    "Juridische dienstverlening",
    "Overig",
]

LAND_OPTIES = (
    ["Nederland", "Overig EU-land (laag risico)"]
    + sorted(EU_HOOG_RISICO_LANDEN)
    + ["Overig niet-EU land"]
)

DOEL_OPTIES = [
    "Jaarrekening en belastingaangifte",
    "Advisering bedrijfsovername / fusie",
    "Bedrijfsoprichting / -structurering",
    "Herstructurering onderneming",
    "Financieel advies / vermogensplanning",
    "Compliance advies",
    "Overig (vrij invullen)",
]

TRANSACTIEPROFIEL_OPTIES = [
    "Reguliere bedrijfsactiviteiten, geen bijzondere transacties",
    "Seizoensgebonden omzetpatroon",
    "Incidentele grote transacties",
    "Internationaal betalingsverkeer",
    "Wisselend / onregelmatig patroon",
    "Overig (vrij invullen)",
]

_ID_SOORT = ["Paspoort", "Identiteitskaart", "Rijbewijs", "Verblijfsdocument"]
_WIJZE_ID = ["Fysiek", "Op afstand / video", "Op afstand / kopie ID"]
UBO_GRONDSLAG = {
    "vof": ["Winstverdeling >25%", "Zeggenschap >25%", "Beide >25%", "Pseudo-UBO (alle vennoten/maten)"],
    "cv": ["Kapitaalinbreng >25%", "Winstrecht >25%", "Beide >25%", "Pseudo-UBO — beherend vennoot"],
    "bv": ["Eigendom >25%", "Zeggenschap >25%", "Eigendom én zeggenschap >25%", "Pseudo-UBO (geen UBO gevonden)"],
    "stichting": [">25% zeggenschap", ">25% begunstiging bij uitkering", "Geen UBO (breed beneficiariaat / ANBI)", "Pseudo-UBO — volledig bestuur"],
    "vereniging": [">25% zeggenschap via stemrecht", "Geen UBO (gespreid stemrecht)", "Pseudo-UBO — volledig bestuur"],
}

INTEGRITEITSVRAGEN = [
    {
        "nr": 4,
        "vraag": "Is er aanleiding integriteitsrisico's te onderkennen op basis van de identiteit en (zakelijke) reputatie van de cliënt, (pseudo-)UBO of vertegenwoordiger?",
        "type": "ja_nee",
        "kritiek": False,
    },
    {
        "nr": 5,
        "vraag": "Is er aanleiding integriteitsrisico's te onderkennen op basis van de houding van de cliënt, (pseudo-)UBO of vertegenwoordiger (gelet op NBA-nadere voorschriften)?",
        "type": "ja_nee",
        "kritiek": False,
    },
    {
        "nr": 6,
        "vraag": "Is er aanleiding integriteitsrisico's te onderkennen op basis van de aard van de werkzaamheden van de cliënt (inclusief bedrijfsvoering)?",
        "type": "ja_nee",
        "kritiek": False,
    },
    {
        "nr": 7,
        "vraag": "Zijn er transacties met verbonden partijen die aanleiding geven tot integriteitsrisico's?",
        "type": "ja_nee",
        "kritiek": False,
    },
    {
        "nr": 8,
        "vraag": "Is er aanleiding integriteitsrisico's te onderkennen op basis van de aard van de bedrijfstak?",
        "type": "ja_nee",
        "kritiek": False,
    },
    {
        "nr": 9,
        "vraag": "Reden accountantswissel / waarom gekozen voor ons kantoor?",
        "type": "vrij_tekst",
        "kritiek": False,
    },
    {
        "nr": 10,
        "vraag": "Is er sprake van onredelijke druk van de kant van de cliënt op ons kantoor?",
        "type": "ja_nee",
        "kritiek": False,
    },
    {
        "nr": 11,
        "vraag": "Is er informatie bekend die erop wijst dat de cliënt zich bezighoudt met witwassen, financiering van terrorisme of andere criminele activiteiten?",
        "type": "ja_nee",
        "kritiek": True,
    },
]

RISICO_KLEUR = {"HOOG": "🔴", "MIDDEN": "🟡", "LAAG": "🟢"}

# ──────────────────────────────────────────────
# KvK API
# ──────────────────────────────────────────────

def zoek_kvk(zoekterm: str, op_naam: bool = False) -> dict:
    if DEMO_MODUS:
        return {"success": True, "resultaten": [DEMO_KVK_DATA]}
    try:
        param = "naam" if op_naam else "kvkNummer"
        resp = _kvk_session.get(f"{KVK_BASE_URL}/zoeken", params={param: zoekterm}, timeout=10)
        resp.raise_for_status()
        resultaten = [r for r in resp.json().get("resultaten", []) if r.get("type") != "rechtspersoon"]
        if resultaten:
            return {"success": True, "resultaten": resultaten}
        return {"success": False, "error": "Geen resultaten gevonden."}
    except requests.exceptions.HTTPError as e:
        if e.response.status_code == 404:
            return {"success": False, "niet_gevonden": True, "error": "Niet gevonden"}
        return {"success": False, "error": f"KvK-service geeft een fout ({e.response.status_code}). Probeer het opnieuw."}
    except requests.exceptions.RequestException as e:
        return {"success": False, "error": str(e)}


def haal_basisprofiel(kvk_nummer: str, href: str = None) -> dict:
    if DEMO_MODUS:
        return {"success": True, "data": DEMO_BASISPROFIEL}
    url = href or f"https://api.kvk.nl/api/v1/basisprofielen/{kvk_nummer}"
    try:
        resp = _kvk_session.get(url, timeout=10)
        resp.raise_for_status()
        return {"success": True, "data": resp.json()}
    except requests.exceptions.RequestException as e:
        return {"success": False, "error": str(e)}



# ──────────────────────────────────────────────
# OpenSanctions
# ──────────────────────────────────────────────

def screen_opensanctions(naam: str, schema: str = "Company") -> dict:
    if not OPENSANCTIONS_API_KEY:
        return {"success": False, "geen_key": True, "error": "Geen OpenSanctions API key."}
    try:
        payload = {"queries": {"entity": {"schema": schema, "properties": {"name": [naam]}}}}
        resp = requests.post(
            "https://api.opensanctions.org/match/default",
            headers={"Authorization": f"ApiKey {OPENSANCTIONS_API_KEY}"},
            json=payload,
            timeout=15,
        )
        resp.raise_for_status()
        resultaten = resp.json().get("responses", {}).get("entity", {}).get("results", [])
        hits = [r for r in resultaten if r.get("score", 0) >= 0.70]
        hoog = [r for r in resultaten if r.get("score", 0) >= 0.85]
        return {"success": True, "resultaten": resultaten[:10], "hits": len(hits), "hoog_risico": len(hoog)}
    except requests.exceptions.RequestException as e:
        return {"success": False, "error": str(e)}

# ──────────────────────────────────────────────
# Adverse media
# ──────────────────────────────────────────────

def zoek_adverse_media(naam: str) -> dict:
    naam_lower = naam.lower()
    zoektermen = [
        f'"{naam}" fraude', f'"{naam}" oplichting', f'"{naam}" witwassen',
        f'"{naam}" faillissement', f'"{naam}" sanctie',
    ]
    resultaten = []
    gezien = set()
    try:
        with DDGS() as ddgs:
            for term in zoektermen:
                for r in ddgs.text(term, max_results=5, region="nl-nl"):
                    url = r.get("href", "")
                    titel = r.get("title", "")
                    tekst = r.get("body", "")
                    if naam_lower not in titel.lower() and naam_lower not in tekst.lower():
                        continue
                    if url in gezien:
                        continue
                    gezien.add(url)
                    resultaten.append({"titel": titel, "url": url, "tekst": tekst[:250]})
        return {"success": True, "resultaten": resultaten}
    except Exception as e:
        return {"success": False, "error": str(e)}

# ──────────────────────────────────────────────
# Risicobeoordeling
# ──────────────────────────────────────────────

def bereken_risico(basisprofiel, screening_bedrijf, screening_personen, land, integriteit_antwoorden):
    score = 0
    factoren = []

    for sbi in get_sbi_codes(basisprofiel):
        code = str(sbi.get("sbiCode", ""))
        omschrijving = sbi.get("sbiOmschrijving", "")
        if code in HOOG_RISICO_SBI:
            score += 3
            factoren.append(f"Hoog-risico sector: {omschrijving} (SBI {code})")
        elif code in MIDDEN_RISICO_SBI:
            score += 1
            factoren.append(f"Aandachtssector: {omschrijving} (SBI {code})")

    rechtsvorm = get_rechtsvorm(basisprofiel)
    if any(rv in rechtsvorm for rv in HOOG_RISICO_RECHTSVORM):
        score += 1
        factoren.append(f"Verhoogde aandacht rechtsvorm: {rechtsvorm}")

    if land in EU_HOOG_RISICO_LANDEN:
        score += 4
        factoren.append(f"EU/FATF hoog-risico land: {land}")
    elif land not in ("Nederland", "Overig EU-land (laag risico)"):
        score += 1
        factoren.append(f"Niet-EU land: {land}")

    if screening_bedrijf.get("success"):
        hoog = screening_bedrijf.get("hoog_risico", 0)
        hits = screening_bedrijf.get("hits", 0)
        if hoog > 0:
            score += 10
            factoren.append(f"SANCTIELIJST BEDRIJF: {hoog} hoge-score treffer(s)!")
        elif hits > 0:
            score += 5
            factoren.append(f"Sanctielijst bedrijf: {hits} mogelijke treffer(s)")

    for naam, res in (screening_personen or {}).items():
        if res.get("success"):
            hoog = res.get("hoog_risico", 0)
            hits = res.get("hits", 0)
            if hoog > 0:
                score += 10
                factoren.append(f"SANCTIELIJST PERSOON '{naam}': {hoog} treffer(s)!")
            elif hits > 0:
                score += 5
                factoren.append(f"Sanctielijst persoon '{naam}': {hits} mogelijke treffer(s)")

    for vraag in INTEGRITEITSVRAGEN:
        if vraag["type"] != "ja_nee":
            continue
        antwoord = (integriteit_antwoorden or {}).get(str(vraag["nr"]), {})
        if antwoord.get("waarde") == "Ja":
            if vraag.get("kritiek"):
                score += 10
                factoren.append(f"KRITIEK — vraag {vraag['nr']}: melding compliance officer vereist!")
            else:
                score += 1
                toel = antwoord.get("toelichting", "")
                factoren.append(f"Integriteitsrisico (vraag {vraag['nr']}){': ' + toel[:50] if toel else ''}")

    if score >= 8:
        return "HOOG", score, factoren
    if score >= 3:
        return "MIDDEN", score, factoren
    return "LAAG", score, factoren


def get_cdd_vorm(risico_klasse: str):
    if risico_klasse == "HOOG":
        return "Verscherpt cliëntenonderzoek", "art. 8 WWFT"
    return "Standaard cliëntenonderzoek", "art. 3 WWFT"

# ──────────────────────────────────────────────
# Helpers basisprofiel
# ──────────────────────────────────────────────

def adres_str(kvk_data: dict) -> str:
    adres = kvk_data.get("adres", {}).get("binnenlandsAdres", {})
    delen = [adres.get("straatnaam", ""), str(adres.get("huisnummer", "")),
             adres.get("postcode", ""), adres.get("plaats", "")]
    return " ".join(d for d in delen if d).strip()

def get_rechtsvorm(bp: dict) -> str:
    return bp.get("_embedded", {}).get("eigenaar", {}).get("rechtsvorm", "-")

def get_sbi_codes(bp: dict) -> list:
    return bp.get("sbiActiviteiten", [])

def get_oprichtingsdatum(bp: dict) -> str:
    raw = bp.get("formeleRegistratiedatum") or bp.get("materieleRegistratie", {}).get("datumAanvang", "")
    raw = str(raw) if raw else ""
    if len(raw) == 8 and raw.isdigit():
        return f"{raw[6:8]}-{raw[4:6]}-{raw[0:4]}"
    return raw or "-"

def get_medewerkers(bp: dict) -> str:
    val = bp.get("totaalWerkzamePersonen")
    return str(val) if val is not None else "-"

def get_websites(bp: dict) -> list:
    return bp.get("_embedded", {}).get("hoofdvestiging", {}).get("websites", [])

def get_volledig_adres(bp: dict) -> str:
    adressen = bp.get("_embedded", {}).get("hoofdvestiging", {}).get("adressen", [])
    return adressen[0].get("volledigAdres", "") if adressen else ""

def get_handelsnamen(bp: dict) -> list:
    return [h.get("naam", "") for h in bp.get("handelsnamen", []) if h.get("naam")]

def get_rechtsvorm_route(bp: dict) -> str:
    rv = get_rechtsvorm(bp).lower()
    if "eenmanszaak" in rv:
        return "eenmanszaak"
    if "commanditaire" in rv:
        return "cv"
    if any(x in rv for x in ["vennootschap onder firma", "maatschap"]):
        return "vof"
    if "stichting" in rv:
        return "stichting"
    if "vereniging" in rv:
        return "vereniging"
    return "bv"

def _id_df(met_bevoegdheid: bool):
    d = {"Naam": [""]}
    if met_bevoegdheid:
        d["Bevoegdheid"] = ["Alleen bevoegd"]
    d.update({
        "Geboortedatum": [None], "Geboorteplaats": [""], "Woonplaats": [""],
        "Soort ID": ["Paspoort"], "Documentnummer": [""],
        "Datum uitgifte": [None], "Plaats uitgifte": [""],
        "Geldig tot": [None], "Datum gezien": [None],
        "Wijze identificatie": ["Fysiek"],
    })
    return pd.DataFrame(d)

def _id_col_cfg(met_bevoegdheid: bool) -> dict:
    cfg = {"Naam": st.column_config.TextColumn("Naam", width="medium")}
    if met_bevoegdheid:
        cfg["Bevoegdheid"] = st.column_config.SelectboxColumn(
            "Bevoegdheid", options=["Alleen bevoegd", "Gezamenlijk bevoegd"], width="medium"
        )
    cfg.update({
        "Geboortedatum": st.column_config.DateColumn("Geboortedatum", width="small"),
        "Geboorteplaats": st.column_config.TextColumn("Geboorteplaats", width="small"),
        "Woonplaats": st.column_config.TextColumn("Woonplaats", width="small"),
        "Soort ID": st.column_config.SelectboxColumn("Soort ID", options=_ID_SOORT, width="small"),
        "Documentnummer": st.column_config.TextColumn("Documentnummer", width="small"),
        "Datum uitgifte": st.column_config.DateColumn("Datum uitgifte", width="small"),
        "Plaats uitgifte": st.column_config.TextColumn("Plaats uitgifte", width="small"),
        "Geldig tot": st.column_config.DateColumn("Geldig tot", width="small"),
        "Datum gezien": st.column_config.DateColumn("Datum gezien", width="small"),
        "Wijze identificatie": st.column_config.SelectboxColumn(
            "Wijze identificatie", options=_WIJZE_ID, width="small"
        ),
    })
    return cfg

# ──────────────────────────────────────────────
# Word generatie
# ──────────────────────────────────────────────

def _rij(tabel, label: str, waarde):
    rij = tabel.add_row()
    rij.cells[0].text = label
    rij.cells[1].text = str(waarde) if waarde is not None else "-"
    for para in rij.cells[0].paragraphs:
        for run in para.runs:
            run.bold = True


def genereer_word(kvk_data, basisprofiel, personen_data,
                  screening_bedrijf, screening_personen, media,
                  integriteit_antwoorden, risico_klasse, score, factoren,
                  cdd_vorm, cdd_artikel, doel_aard, toelichting,
                  medewerker, referentie, nu) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    h = doc.add_heading("WWFT Cliëntenonderzoek", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(kvk_data.get("naam", ""))
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph("")

    # 1. Rapportgegevens
    doc.add_heading("1. Rapportgegevens", 1)
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    _rij(t, "Referentie", referentie)
    _rij(t, "Datum / tijd", nu.strftime("%d-%m-%Y %H:%M"))
    _rij(t, "Uitvoerder", medewerker or "[invullen]")
    _rij(t, "Beoordeeld door", "")
    _rij(t, "Datum beoordeling", "")
    doc.add_paragraph("")

    # 2. Cliëntgegevens
    doc.add_heading("2. Cliëntgegevens", 1)
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    _rij(t, "KvK-nummer", kvk_data.get("kvkNummer", ""))
    _rij(t, "Naam", kvk_data.get("naam", ""))
    _rij(t, "Rechtsvorm", get_rechtsvorm(basisprofiel))
    _rij(t, "Adres", get_volledig_adres(basisprofiel) or adres_str(kvk_data))
    _rij(t, "Oprichtingsdatum", get_oprichtingsdatum(basisprofiel))
    _rij(t, "Medewerkers", get_medewerkers(basisprofiel))
    handelsnamen = get_handelsnamen(basisprofiel)
    if len(handelsnamen) > 1:
        _rij(t, "Handelsnamen", ", ".join(handelsnamen))
    websites = get_websites(basisprofiel)
    if websites:
        _rij(t, "Website", ", ".join(websites))
    doc.add_paragraph("")

    sbi_codes = get_sbi_codes(basisprofiel)
    if sbi_codes:
        doc.add_heading("2a. Activiteiten (SBI-codes)", 2)
        t = doc.add_table(rows=0, cols=3)
        t.style = "Table Grid"
        hr = t.add_row()
        for i, h_tekst in enumerate(["SBI-code", "Omschrijving", "Risico"]):
            p = hr.cells[i].paragraphs[0]
            run = p.add_run(h_tekst)
            run.bold = True
        for sbi in sbi_codes:
            code = str(sbi.get("sbiCode", ""))
            omschrijving = sbi.get("sbiOmschrijving", "")
            risico = "Hoog" if code in HOOG_RISICO_SBI else "Aandacht" if code in MIDDEN_RISICO_SBI else "Standaard"
            rij = t.add_row()
            rij.cells[0].text = code
            rij.cells[1].text = omschrijving
            rij.cells[2].text = risico
        doc.add_paragraph("")

    # 3. Doel en aard zakelijke relatie
    doc.add_heading("3. Doel en aard zakelijke relatie (art. 3 lid 2 sub b WWFT)", 1)
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    _rij(t, "Type dienstverlening", doel_aard.get("dienstverlening", ""))
    _rij(t, "Land cliënt / UBO's", doel_aard.get("land", ""))
    _rij(t, "Doel zakelijke relatie", doel_aard.get("doel", ""))
    _rij(t, "Verwacht transactieprofiel", doel_aard.get("transactieprofiel", ""))
    doc.add_paragraph("")

    # 4. Personen — route-afhankelijk
    route_w = personen_data.get("route", "bv")
    vertegens_w = personen_data.get("vertegenwoordigers", [])
    ubos_w = personen_data.get("ubos", [])

    def _persoon_tabel(doc, vertegens, toon_bevoegdheid=True):
        for v in vertegens:
            t = doc.add_table(rows=0, cols=2)
            t.style = "Table Grid"
            _rij(t, "Naam", v.get("naam", ""))
            if toon_bevoegdheid and v.get("bevoegdheid"):
                _rij(t, "Bevoegdheid", v["bevoegdheid"])
            _rij(t, "Geboortedatum", v.get("geboortedatum", "") or "-")
            _rij(t, "Geboorteplaats", v.get("geboorteplaats", "") or "-")
            _rij(t, "Woonplaats", v.get("woonplaats", "") or "-")
            _rij(t, "Soort identiteitsdocument", v.get("soort_id", "") or "-")
            _rij(t, "Documentnummer", v.get("documentnummer", "") or "-")
            _rij(t, "Datum van uitgifte", v.get("datum_uitgifte", "") or "-")
            _rij(t, "Plaats van uitgifte", v.get("plaats_uitgifte", "") or "-")
            _rij(t, "Geldig tot", v.get("geldig_tot", "") or "-")
            _rij(t, "Datum gezien", v.get("datum_gezien", "") or "-")
            _rij(t, "Wijze van identificatie", v.get("wijze_identificatie", "") or "-")
            doc.add_paragraph("")

    def _ubo_tabel(doc, ubos):
        if ubos:
            met_rol = any(u.get("rol") for u in ubos)
            headers = ["Naam", "Rol", "Grondslag", "Belang / toelichting"] if met_rol else ["Naam", "Grondslag", "Belang / toelichting"]
            t = doc.add_table(rows=0, cols=len(headers))
            t.style = "Table Grid"
            hr = t.add_row()
            for i, h_tekst in enumerate(headers):
                hr.cells[i].paragraphs[0].add_run(h_tekst).bold = True
            for u in ubos:
                rij = t.add_row()
                if met_rol:
                    rij.cells[0].text = u.get("naam", "")
                    rij.cells[1].text = u.get("rol", "")
                    rij.cells[2].text = u.get("grondslag", "")
                    rij.cells[3].text = u.get("belang", "")
                else:
                    rij.cells[0].text = u.get("naam", "")
                    rij.cells[1].text = u.get("grondslag", "")
                    rij.cells[2].text = u.get("belang", "")
        else:
            doc.add_paragraph("Geen UBO's geregistreerd.")
        doc.add_paragraph("")

    if route_w == "eenmanszaak":
        doc.add_heading("4. Identificatie ondernemer (art. 33 WWFT)", 1)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        _rij(t, "KvK-uittreksel datum", personen_data.get("kvk_uittreksel_datum", "") or "-")
        _rij(t, "Opmerking", "Eenmanszaak — ondernemer is cliënt en uiteindelijk belanghebbende")
        doc.add_paragraph("")
        if vertegens_w:
            _persoon_tabel(doc, vertegens_w, toon_bevoegdheid=False)
        else:
            doc.add_paragraph("Geen ondernemer geregistreerd.")
            doc.add_paragraph("")

    elif route_w == "vof":
        doc.add_heading("4. Identificatie vennoten (art. 33 WWFT)", 1)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        _rij(t, "KvK-uittreksel datum", personen_data.get("kvk_uittreksel_datum", "") or "-")
        _rij(t, "Bevoegdheid gecontroleerd",
             "Ja" if personen_data.get("bevoegdheid_gecontroleerd") else "Nee")
        _rij(t, "Vennootschapsovereenkomst geraadpleegd",
             "Ja" if personen_data.get("overeenkomst_geraadpleegd") else "Nee")
        doc.add_paragraph("")
        if vertegens_w:
            _persoon_tabel(doc, vertegens_w, toon_bevoegdheid=True)
        else:
            doc.add_paragraph("Geen vennoot geregistreerd.")
            doc.add_paragraph("")

        doc.add_heading("4a. UBO-vaststelling (winst-/zeggenschapsverdeling, art. 3 lid 2 sub c WWFT)", 2)
        _ubo_tabel(doc, ubos_w)

        doc.add_heading("4b. UBO-register controle", 2)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        _rij(t, "UBO-register gecontroleerd (kvk.nl)",
             "Ja" if personen_data.get("ubo_register_gecontroleerd") else "Nee")
        _rij(t, "Datum controle", personen_data.get("ubo_register_datum", "") or "-")
        doc.add_paragraph("")

    elif route_w == "cv":
        doc.add_heading("4. Identificatie beherend vennoot (art. 33 WWFT)", 1)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        _rij(t, "KvK-uittreksel datum", personen_data.get("kvk_uittreksel_datum", "") or "-")
        _rij(t, "Bevoegdheid gecontroleerd",
             "Ja" if personen_data.get("bevoegdheid_gecontroleerd") else "Nee")
        _rij(t, "CV-overeenkomst geraadpleegd",
             "Ja" if personen_data.get("overeenkomst_geraadpleegd") else "Nee")
        doc.add_paragraph("")
        if vertegens_w:
            _persoon_tabel(doc, vertegens_w, toon_bevoegdheid=True)
        else:
            doc.add_paragraph("Geen beherend vennoot geregistreerd.")
            doc.add_paragraph("")

        doc.add_heading("4a. UBO-vaststelling beherend én stille vennoten (art. 3 lid 2 sub c WWFT)", 2)
        doc.add_paragraph("Stille/commanditaire vennoten tellen mee voor UBO-vaststelling maar treden niet op als vertegenwoordiger.")
        _ubo_tabel(doc, ubos_w)

        doc.add_heading("4b. UBO-register controle", 2)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        _rij(t, "UBO-register gecontroleerd (kvk.nl)",
             "Ja" if personen_data.get("ubo_register_gecontroleerd") else "Nee")
        _rij(t, "Datum controle", personen_data.get("ubo_register_datum", "") or "-")
        doc.add_paragraph("")

    elif route_w in ("stichting", "vereniging"):
        _rvlabel = "Stichting" if route_w == "stichting" else "Vereniging"
        doc.add_heading(f"4. Identificatie bestuurder(s) {_rvlabel} (art. 33 WWFT)", 1)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        _rij(t, "KvK-uittreksel datum", personen_data.get("kvk_uittreksel_datum", "") or "-")
        _rij(t, "Bevoegdheid gecontroleerd",
             "Ja" if personen_data.get("bevoegdheid_gecontroleerd") else "Nee")
        doc.add_paragraph("")
        if vertegens_w:
            _persoon_tabel(doc, vertegens_w, toon_bevoegdheid=True)
        else:
            doc.add_paragraph("Geen bestuurder geregistreerd.")
            doc.add_paragraph("")

        doc.add_heading("4a. UBO-vaststelling (art. 3 lid 2 sub c WWFT)", 2)
        _ubo_tabel(doc, ubos_w)

        doc.add_heading("4b. UBO-register controle", 2)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        _rij(t, "UBO-register gecontroleerd (kvk.nl)",
             "Ja" if personen_data.get("ubo_register_gecontroleerd") else "Nee")
        _rij(t, "Datum controle", personen_data.get("ubo_register_datum", "") or "-")
        doc.add_paragraph("")

    else:  # bv / nv / overig
        doc.add_heading("4. Identificatie vertegenwoordiger(s) (art. 33 WWFT)", 1)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        _rij(t, "KvK-uittreksel datum", personen_data.get("kvk_uittreksel_datum", "") or "-")
        _rij(t, "Bevoegdheid gecontroleerd",
             "Ja" if personen_data.get("bevoegdheid_gecontroleerd") else "Nee")
        doc.add_paragraph("")
        if vertegens_w:
            _persoon_tabel(doc, vertegens_w, toon_bevoegdheid=True)
        else:
            doc.add_paragraph("Geen vertegenwoordiger geregistreerd.")
            doc.add_paragraph("")

        doc.add_heading("4a. UBO-vaststelling (art. 3 lid 2 sub c WWFT)", 2)
        _ubo_tabel(doc, ubos_w)

        doc.add_heading("4b. UBO-register controle", 2)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        _rij(t, "UBO-register gecontroleerd (kvk.nl)",
             "Ja" if personen_data.get("ubo_register_gecontroleerd") else "Nee")
        _rij(t, "Datum controle", personen_data.get("ubo_register_datum", "") or "-")
        doc.add_paragraph("")

    # 5. Persoonsscreening
    doc.add_heading("5. Persoonsscreening (OpenSanctions — PEP & sancties)", 1)
    if screening_personen:
        t = doc.add_table(rows=0, cols=3)
        t.style = "Table Grid"
        hr = t.add_row()
        for i, h_tekst in enumerate(["Naam", "Uitkomst", "Treffers (>70%)"]):
            hr.cells[i].paragraphs[0].add_run(h_tekst).bold = True
        for naam, res in screening_personen.items():
            rij = t.add_row()
            rij.cells[0].text = naam
            if not res.get("success"):
                rij.cells[1].text = "Niet uitgevoerd" if res.get("geen_key") else "Fout"
                rij.cells[2].text = "-"
            else:
                hoog = res.get("hoog_risico", 0)
                hits = res.get("hits", 0)
                if hoog > 0:
                    rij.cells[1].text = f"HOOG RISICO ({hoog} treffers)"
                elif hits > 0:
                    rij.cells[1].text = f"Mogelijke treffers ({hits})"
                else:
                    rij.cells[1].text = "Geen treffers"
                rij.cells[2].text = str(hits)
    else:
        doc.add_paragraph("Geen persoonsscreening uitgevoerd.")
    doc.add_paragraph("")

    # 6. Sanctiescreening bedrijf
    doc.add_heading("6. Sanctie- en PEP-screening bedrijf (OpenSanctions)", 1)
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    _rij(t, "Bron", "OpenSanctions (EU, VN, OFAC, ~40 bronnen)")
    if not screening_bedrijf.get("success"):
        _rij(t, "Uitkomst", "Niet uitgevoerd" if screening_bedrijf.get("geen_key") else "Mislukt")
    else:
        hoog = screening_bedrijf.get("hoog_risico", 0)
        hits = screening_bedrijf.get("hits", 0)
        if hoog > 0:
            _rij(t, "Uitkomst", f"HOGE SCORE TREFFERS ({hoog}) — direct actie vereist")
        elif hits > 0:
            _rij(t, "Uitkomst", f"Mogelijke treffers ({hits}) — nader onderzoek")
        else:
            _rij(t, "Uitkomst", "Geen treffers gevonden")
        _rij(t, "Treffers (>70%)", str(screening_bedrijf.get("hits", 0)))
        _rij(t, "Hoge score (>85%)", str(hoog))
    doc.add_paragraph("")

    # 7. Adverse media
    doc.add_heading("7. Adverse media search", 1)
    if not media.get("success"):
        doc.add_paragraph("Niet uitgevoerd of mislukt.")
    else:
        media_resultaten = media.get("resultaten", [])
        if not media_resultaten:
            doc.add_paragraph("Geen negatief nieuws gevonden waarbij de bedrijfsnaam voorkomt.")
        else:
            doc.add_paragraph(f"{len(media_resultaten)} relevante resultaten gevonden — beoordeel handmatig:")
            for r in media_resultaten:
                p_item = doc.add_paragraph(style="List Bullet")
                p_item.add_run(r.get("titel", "")[:120]).bold = True
                doc.add_paragraph(r.get("tekst", "")[:250])
                doc.add_paragraph(r.get("url", ""))
    doc.add_paragraph("")

    # 8. Integriteitsvragen
    doc.add_heading("8. Integriteitsvragen (NBA NV WWFT)", 1)
    t = doc.add_table(rows=0, cols=3)
    t.style = "Table Grid"
    hr = t.add_row()
    for i, h_tekst in enumerate(["Vraag", "Antwoord", "Toelichting"]):
        hr.cells[i].paragraphs[0].add_run(h_tekst).bold = True
    for vraag in INTEGRITEITSVRAGEN:
        antwoord = (integriteit_antwoorden or {}).get(str(vraag["nr"]), {})
        rij = t.add_row()
        rij.cells[0].text = f"{vraag['nr']}. {vraag['vraag']}"
        if vraag["type"] == "vrij_tekst":
            rij.cells[1].text = "—"
            rij.cells[2].text = antwoord.get("tekst", "")
        else:
            rij.cells[1].text = antwoord.get("waarde", "-")
            rij.cells[2].text = antwoord.get("toelichting", "")
    doc.add_paragraph("")

    # 9. Risicobeoordeling
    doc.add_heading("9. Risicobeoordeling", 1)
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    _rij(t, "Risicoclassificatie", risico_klasse)
    _rij(t, "Totaalscore", str(score))
    _rij(t, "Vereiste CDD-vorm", f"{cdd_vorm} ({cdd_artikel})")
    factoren_tekst = "\n".join(f"- {f}" for f in factoren) if factoren else "Geen bijzondere factoren"
    _rij(t, "Risicobepalende factoren", factoren_tekst)
    _rij(t, "Toelichting medewerker", toelichting.strip() if toelichting else "")
    doc.add_paragraph("")

    # 10. Conclusie
    doc.add_heading("10. Conclusie en vervolgactie", 1)
    if risico_klasse == "HOOG":
        conclusie = "Verscherpt cliëntenonderzoek vereist conform art. 8 WWFT. Melding compliance officer verplicht."
    elif risico_klasse == "MIDDEN":
        conclusie = "Standaard cliëntenonderzoek van toepassing. Verhoogde alertheid gedurende de relatie aanbevolen."
    else:
        conclusie = "Standaard cliëntenonderzoek van toepassing."
    doc.add_paragraph(conclusie)
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    _rij(t, "Volgende hertoetsing", nu.strftime("%d-%m-") + str(nu.year + 1))
    doc.add_paragraph("")

    # 11. Akkoordverklaring
    doc.add_heading("11. Akkoordverklaring", 1)
    t = doc.add_table(rows=0, cols=4)
    t.style = "Table Grid"
    hr = t.add_row()
    for i, h_tekst in enumerate(["Rol", "Naam", "Datum", "Paraaf"]):
        hr.cells[i].paragraphs[0].add_run(h_tekst).bold = True
    for rol, naam_val in [("Uitvoerder", medewerker), ("Compliance review", "")]:
        rij = t.add_row()
        rij.cells[0].text = rol
        rij.cells[1].text = naam_val or ""
        rij.cells[2].text = ""
        rij.cells[3].text = ""
    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run(f"Gegenereerd door WWFT Check Tool  —  {referentie}  —  {nu.strftime('%d-%m-%Y %H:%M')}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ──────────────────────────────────────────────
# Sidebar
# ──────────────────────────────────────────────

with st.sidebar:
    st.markdown("## 🔍 WWFT Check Tool")
    st.caption("Wet ter voorkoming van witwassen\nen financieren van terrorisme")

    if DEMO_MODUS:
        st.warning("**Demo modus** — voorbeelddata")
    else:
        st.success("KvK API verbonden")

    st.divider()

    _pagina = st.radio(
        "Navigatie",
        ["🔍 Cliëntenonderzoek", "📖 Handleiding"],
        label_visibility="collapsed",
    )

    st.divider()

    # Voortgang
    _done = {
        "kvk":         "kvk_data"          in st.session_state,
        "doel":        "doel_aard"          in st.session_state,
        "personen":    "personen_data"      in st.session_state,
        "screening":   "screening_bedrijf"  in st.session_state,
        "integriteit": "integriteitsvragen" in st.session_state,
    }
    st.markdown("**Voortgang**")
    _stappen_ui = [
        ("Bedrijf zoeken",     "kvk"),
        ("Doel en aard",       "doel"),
        ("Personen & UBO's",   "personen"),
        ("Screening",          "screening"),
        ("Integriteitsvragen", "integriteit"),
        ("Risicobeoordeling",  "integriteit"),
        ("Rapport downloaden", "integriteit"),
    ]
    _actief_geweest = False
    for _lbl, _key in _stappen_ui:
        if _done.get(_key):
            st.markdown(f'<div class="prog-item prog-done">✅ {_lbl}</div>', unsafe_allow_html=True)
        elif not _actief_geweest:
            st.markdown(f'<div class="prog-item prog-active">▶ {_lbl}</div>', unsafe_allow_html=True)
            _actief_geweest = True
        else:
            st.markdown(f'<div class="prog-item prog-todo">○ {_lbl}</div>', unsafe_allow_html=True)

    st.divider()
    st.caption("v0.3 – Concept")

# ──────────────────────────────────────────────
# Handleiding-pagina
# ──────────────────────────────────────────────

_APP_DIR = os.path.dirname(os.path.abspath(__file__))

if _pagina == "📖 Handleiding":
    st.title("📖 Handleiding Wwft-cliëntenonderzoek")
    st.caption("Praktische naslaggids voor medewerkers — versie 1")

    tab_infographic, tab_tekst = st.tabs(["🗺️ Interactief stappenplan", "📄 Volledige handleiding"])

    with tab_infographic:
        st.markdown("Klik op een stap in de tijdlijn of gebruik de knoppen om door het proces te navigeren.")
        _html_path = os.path.join(_APP_DIR, "wwft_stappenplan_visueel.html")
        with open(_html_path, encoding="utf-8") as _f:
            _fragment = _f.read()
        _full_html = """<!DOCTYPE html>
<html lang="nl">
<head>
<meta charset="utf-8">
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@3/tabler-icons.min.css">
<style>
:root{--border:#dee2e6;--surface-1:#f1f3f5;--text-muted:#868e96;--text-accent:#378ADD;--text-secondary:#495057}
*{box-sizing:border-box}
body{font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,sans-serif;margin:0;padding:16px 8px;background:#fff;color:#212529}
button{cursor:pointer;padding:8px 18px;border:1px solid #dee2e6;border-radius:8px;background:#fff;font-size:14px;transition:background .15s}
button:hover{background:#f1f3f5}
</style>
</head>
<body>
<script>function sendPrompt(){}</script>
""" + _fragment + "\n</body>\n</html>"
        components.html(_full_html, height=430, scrolling=False)

    with tab_tekst:
        _md_path = os.path.join(_APP_DIR, "wwft-handleiding-medewerkers_v1.md")
        with open(_md_path, encoding="utf-8") as _f:
            _md_content = _f.read()
        st.markdown(_md_content)

    st.stop()

# ──────────────────────────────────────────────
# Stap 1: Zoeken
# ──────────────────────────────────────────────

st.markdown("## Cliëntenonderzoek starten")
st.caption("Zoek het bedrijf op in het Handelsregister om het onderzoek te starten.")

with st.form("kvk_form"):
    col1, col2, col3 = st.columns([2, 1, 1])
    with col1:
        zoekterm_input = st.text_input(
            "Zoeken op naam of KvK-nummer",
            placeholder="bijv. Join Administraties B.V.  of  12345678",
        )
    with col2:
        zoektype = st.radio("Zoeken op", ["Naam", "KvK-nummer"], horizontal=True)
    with col3:
        st.write("")
        st.write("")
        zoeken = st.form_submit_button("Zoeken →", use_container_width=True, type="primary")

if zoeken:
    zoekterm = zoekterm_input.strip()
    if not zoekterm:
        st.error("Voer een naam of KvK-nummer in.")
        st.stop()
    with st.spinner("KvK raadplegen…"):
        zoek_result = zoek_kvk(zoekterm, op_naam=(zoektype == "Naam"))
    if not zoek_result["success"]:
        if zoek_result.get("niet_gevonden"):
            st.warning(f"Geen bedrijf gevonden voor **{zoekterm}**. Controleer de schrijfwijze of het KvK-nummer en probeer opnieuw.")
        else:
            st.error(zoek_result["error"])
        st.stop()
    for k in ["kvk_data", "basisprofiel", "functionarissen", "doel_aard",
              "personen_data", "screening_bedrijf", "screening_personen",
              "media", "integriteitsvragen"]:
        st.session_state.pop(k, None)
    st.session_state["zoekresultaten"] = zoek_result["resultaten"]
    st.rerun()

if "zoekresultaten" in st.session_state and "kvk_data" not in st.session_state:
    resultaten = st.session_state["zoekresultaten"]
    if len(resultaten) == 1:
        gekozen = resultaten[0]
    else:
        st.subheader(f"🔎 {len(resultaten)} resultaten – kies het juiste bedrijf")
        opties = {
            f"{r.get('naam','?')} — KvK {r.get('kvkNummer','?')} — "
            f"{r.get('adres',{}).get('binnenlandsAdres',{}).get('plaats','')}": r
            for r in resultaten
        }
        keuze_label = st.radio("Selecteer bedrijf", list(opties.keys()))
        if not st.button("Dit bedrijf selecteren →", type="primary"):
            st.stop()
        gekozen = opties[keuze_label]

    links = {lnk["rel"]: lnk["href"] for lnk in gekozen.get("links", [])}
    kvk_nummer = gekozen.get("kvkNummer", "")
    with st.spinner("Basisprofiel ophalen…"):
        profiel_result = haal_basisprofiel(kvk_nummer, href=links.get("basisprofiel"))
    st.session_state["kvk_data"] = gekozen
    st.session_state["basisprofiel"] = profiel_result.get("data", {})
    st.session_state["kvk_nummer"] = kvk_nummer
    st.session_state["functionarissen"] = []
    st.rerun()

# ──────────────────────────────────────────────
# Stap 2: Bedrijfsgegevens
# ──────────────────────────────────────────────

if "kvk_data" not in st.session_state:
    st.stop()

kvk_data = st.session_state["kvk_data"]
basisprofiel = st.session_state["basisprofiel"]
functionarissen = st.session_state.get("functionarissen", [])
bedrijfsnaam = kvk_data.get("naam", "Onbekend")

st.divider()
st.subheader(f"📋 {bedrijfsnaam}")

datum_uit = basisprofiel.get("datumUitschrijving")
if datum_uit:
    st.error(f"⚠️ Dit bedrijf is uitgeschreven uit het handelsregister op {datum_uit}.")

rechtsvorm = get_rechtsvorm(basisprofiel)
volledig_adres = get_volledig_adres(basisprofiel) or adres_str(kvk_data)
handelsnamen = get_handelsnamen(basisprofiel)
websites = get_websites(basisprofiel)
sbi_codes = get_sbi_codes(basisprofiel)

col1, col2, col3 = st.columns(3)
col1.metric("KvK-nummer", kvk_data.get("kvkNummer", "-"))
col1.metric("Rechtsvorm", rechtsvorm)
col2.metric("Adres", volledig_adres or "-")
col2.metric("Oprichtingsdatum", get_oprichtingsdatum(basisprofiel))
col3.metric("Medewerkers", get_medewerkers(basisprofiel))

if len(handelsnamen) > 1:
    st.info(f"**Handelsnamen:** {', '.join(handelsnamen)}")
if websites:
    st.write(f"**Website:** {', '.join(websites)}")

if sbi_codes:
    st.write("**Activiteiten (SBI-codes):**")
    for sbi in sbi_codes:
        code = str(sbi.get("sbiCode", ""))
        omschrijving = sbi.get("sbiOmschrijving", "")
        label = "🔴 Hoog-risico" if code in HOOG_RISICO_SBI else \
                "🟡 Aandacht" if code in MIDDEN_RISICO_SBI else "🟢 Standaard"
        st.write(f"- **{code}**: {omschrijving} — {label}")


# ──────────────────────────────────────────────
# Stap 3: Doel en aard + medewerker
# ──────────────────────────────────────────────

st.divider()
st.subheader("📝 Stap 3: Doel en aard zakelijke relatie")
st.caption("Verplicht vast te stellen conform art. 3 lid 2 sub b WWFT")

if "doel_aard" not in st.session_state:
    with st.form("doel_aard_form"):
        col1, col2 = st.columns(2)
        with col1:
            medewerker_input = st.text_input("Naam uitvoerder", placeholder="Uw naam")
            dienstverlening = st.selectbox("Type dienstverlening", DIENSTVERLENING_OPTIES)
            land = st.selectbox("Land van vestiging / herkomst cliënt en UBO's", LAND_OPTIES)
        with col2:
            doel_keuze = st.selectbox("Doel van de zakelijke relatie", DOEL_OPTIES)
            doel_vrij = ""
            if doel_keuze == "Overig (vrij invullen)":
                doel_vrij = st.text_area("Toelichting doel (verplicht)", height=68)
            transactieprofiel_keuze = st.selectbox("Verwacht transactieprofiel", TRANSACTIEPROFIEL_OPTIES)
            transactieprofiel_vrij = ""
            if transactieprofiel_keuze == "Overig (vrij invullen)":
                transactieprofiel_vrij = st.text_area("Toelichting transactieprofiel", height=68)
        bevestig = st.form_submit_button("Bevestigen en doorgaan →", type="primary")

    if bevestig:
        doel = doel_vrij.strip() if doel_keuze == "Overig (vrij invullen)" else doel_keuze
        tp = transactieprofiel_vrij.strip() if transactieprofiel_keuze == "Overig (vrij invullen)" else transactieprofiel_keuze
        if doel_keuze == "Overig (vrij invullen)" and not doel:
            st.error("Vul het doel van de zakelijke relatie in.")
        else:
            st.session_state["doel_aard"] = {
                "medewerker": medewerker_input.strip(),
                "dienstverlening": dienstverlening,
                "land": land,
                "doel": doel,
                "transactieprofiel": tp,
            }
            st.rerun()
    st.stop()
else:
    da = st.session_state["doel_aard"]
    col1, col2 = st.columns(2)
    col1.write(f"**Uitvoerder:** {da.get('medewerker') or '[niet ingevuld]'}")
    col1.write(f"**Dienstverlening:** {da['dienstverlening']}")
    col1.write(f"**Land:** {da['land']}")
    col2.write(f"**Doel:** {da['doel']}")
    if da.get("transactieprofiel"):
        col2.write(f"**Transactieprofiel:** {da['transactieprofiel']}")
    if st.button("Wijzigen", key="wijzig_doel"):
        for k in ["doel_aard", "personen_data", "screening_bedrijf",
                  "screening_personen", "media", "integriteitsvragen"]:
            st.session_state.pop(k, None)
        st.rerun()

# ──────────────────────────────────────────────
# Stap 4: Personen en UBO-onderzoek
# ──────────────────────────────────────────────

_ROUTE_LABEL = {
    "eenmanszaak": "Identificatie ondernemer",
    "vof": "Identificatie vennoten en UBO-vaststelling",
    "cv": "Identificatie beherend vennoot en UBO-vaststelling",
    "bv": "Identificatie vertegenwoordiger(s) en UBO-vaststelling",
    "stichting": "Identificatie bestuurder(s) en UBO-vaststelling",
    "vereniging": "Identificatie bestuurder(s) en UBO-vaststelling",
}
_VERTEGEN_INFO = {
    "vof": ("Optredende vennoot/maat", "De vennoot of maat die namens de VOF/maatschap optreedt."),
    "cv": ("Beherend vennoot", "De beherend vennoot voert het bestuur en is tekeningsbevoegd. Stille vennoten treden NIET op als vertegenwoordiger maar tellen WEL mee voor de UBO-vaststelling."),
    "bv": ("Vertegenwoordiger(s) — bestuurder(s)", "Bij bestuur door een BV: doorpakken tot een natuurlijk persoon."),
    "stichting": ("Bestuurder(s)", "De bestuurder(s) die namens de stichting optreedt/optreden."),
    "vereniging": ("Bestuurder(s)", "De bestuurder(s) die namens de vereniging optreedt/optreden."),
}
_UBO_INFO = {
    "vof": "Vennoot/maat met >25% economisch belang of zeggenschap. Stel de verdeling vast via de vennootschapsovereenkomst. Niemand boven 25%? Alle vennoten/maten zijn pseudo-UBO.",
    "cv": "Elke vennoot (beherend én stil) met >25% kapitaalinbreng of winstrecht. Let op: stille vennoten worden vaak vergeten — controleer ze apart. Geen UBO? Beherend vennoten zijn pseudo-UBO.",
    "bv": "Aandeelhouder met >25% eigendom of zeggenschap. Bestuurder ≠ automatisch UBO. Holdingstructuur? Doorpakken naar de uiteindelijke natuurlijk persoon.",
    "stichting": "Persoon met >25% zeggenschap of begunstigde van >25% van het vermogen bij uitkering. Bij breed/onbepaald beneficiariaat (ANBI): selecteer 'Geen UBO'. Pseudo-UBO = volledig bestuur.",
    "vereniging": "Lid met >25% stemrecht. Bij gespreid stemrecht: selecteer 'Geen UBO'. Pseudo-UBO = volledig bestuur.",
}

route = get_rechtsvorm_route(basisprofiel)
st.divider()
st.subheader(f"👥 Stap 4: {_ROUTE_LABEL.get(route, 'Personen en UBO-onderzoek')}")

if "personen_data" not in st.session_state:
    persoon_df = None
    ubo_df = None
    kvk_datum = None
    bevoegdheid_check = False
    overeenkomst_check = False
    ubo_register = False
    ubo_datum_input = None

    with st.form("personen_form"):

        # ══ Eenmanszaak ═══════════════════════════════════════════════════
        if route == "eenmanszaak":
            st.info("Eenmanszaak: de ondernemer **is** de cliënt én de uiteindelijk belanghebbende. Geen aparte UBO-vaststelling nodig.")
            st.write("**A. Identificatie onderneming**")
            col1, _ = st.columns(2)
            kvk_datum = col1.date_input("KvK-uittreksel datum opgevraagd", value=None)
            st.divider()
            st.write("**B. Identificatie ondernemer** (art. 33 WWFT)")
            st.caption("Géén kopie met BSN/pasfoto bewaren (AVG). Vastleggen van de brongegevens is voldoende.")
            persoon_df = st.data_editor(
                _id_df(False), column_config=_id_col_cfg(False),
                num_rows="dynamic", use_container_width=True, key="persoon_editor",
            )

        # ══ VOF / Maatschap ════════════════════════════════════════════════
        elif route == "vof":
            st.write("**A. Identificatie vennootschap**")
            col1, col2 = st.columns(2)
            kvk_datum = col1.date_input("KvK-uittreksel datum opgevraagd", value=None)
            bevoegdheid_check = col2.checkbox("Bevoegdheid gecontroleerd in Handelsregister")
            overeenkomst_check = st.checkbox("Vennootschapsovereenkomst geraadpleegd voor winst-/zeggenschapsverdeling")
            st.divider()
            lbl, cap = _VERTEGEN_INFO["vof"]
            st.write(f"**B. {lbl}** (art. 33 WWFT)")
            st.caption(cap)
            persoon_df = st.data_editor(
                _id_df(True), column_config=_id_col_cfg(True),
                num_rows="dynamic", use_container_width=True, key="persoon_editor",
            )
            st.divider()
            st.write("**C. UBO-vaststelling**")
            st.caption(_UBO_INFO["vof"])
            ubo_df = st.data_editor(
                pd.DataFrame({"Naam": [""], "Grondslag": [UBO_GRONDSLAG["vof"][0]], "Belang / toelichting": [""]}),
                column_config={
                    "Naam": st.column_config.TextColumn("Naam", width="medium"),
                    "Grondslag": st.column_config.SelectboxColumn("Grondslag", options=UBO_GRONDSLAG["vof"], width="large"),
                    "Belang / toelichting": st.column_config.TextColumn("Belang / toelichting", width="medium"),
                },
                num_rows="dynamic", use_container_width=True, key="ubo_editor",
            )
            st.divider()
            st.write("**D. UBO-register controle (kvk.nl/ubo-register)**")
            col1, col2 = st.columns(2)
            ubo_register = col1.checkbox("UBO-register gecontroleerd")
            ubo_datum_input = col2.date_input("Datum controle", value=None, key="ubo_reg_datum")

        # ══ CV ═════════════════════════════════════════════════════════════
        elif route == "cv":
            st.write("**A. Identificatie CV**")
            col1, col2 = st.columns(2)
            kvk_datum = col1.date_input("KvK-uittreksel datum opgevraagd", value=None)
            bevoegdheid_check = col2.checkbox("Bevoegdheid gecontroleerd in Handelsregister")
            overeenkomst_check = st.checkbox("CV-overeenkomst geraadpleegd voor kapitaalinbreng en winstverdeling")
            st.divider()
            lbl, cap = _VERTEGEN_INFO["cv"]
            st.write(f"**B. {lbl}** (art. 33 WWFT)")
            st.caption(cap)
            persoon_df = st.data_editor(
                _id_df(True), column_config=_id_col_cfg(True),
                num_rows="dynamic", use_container_width=True, key="persoon_editor",
            )
            st.divider()
            st.write("**C. UBO-vaststelling** (beherend én stille vennoten)")
            st.caption(_UBO_INFO["cv"])
            ubo_df = st.data_editor(
                pd.DataFrame({"Naam": [""], "Rol": ["Beherend vennoot"], "Grondslag": [UBO_GRONDSLAG["cv"][0]], "Belang / toelichting": [""]}),
                column_config={
                    "Naam": st.column_config.TextColumn("Naam", width="medium"),
                    "Rol": st.column_config.SelectboxColumn("Rol", options=["Beherend vennoot", "Stille/commanditaire vennoot"], width="medium"),
                    "Grondslag": st.column_config.SelectboxColumn("Grondslag", options=UBO_GRONDSLAG["cv"], width="large"),
                    "Belang / toelichting": st.column_config.TextColumn("Belang / toelichting", width="medium"),
                },
                num_rows="dynamic", use_container_width=True, key="ubo_editor",
            )
            st.divider()
            st.write("**D. UBO-register controle (kvk.nl/ubo-register)**")
            col1, col2 = st.columns(2)
            ubo_register = col1.checkbox("UBO-register gecontroleerd")
            ubo_datum_input = col2.date_input("Datum controle", value=None, key="ubo_reg_datum")

        # ══ BV / NV / Stichting / Vereniging ══════════════════════════════
        else:
            st.write("**A. Identificatie rechtspersoon**")
            col1, col2 = st.columns(2)
            kvk_datum = col1.date_input("KvK-uittreksel datum opgevraagd", value=None)
            bevoegdheid_check = col2.checkbox("Bevoegdheid bestuurder(s) gecontroleerd in Handelsregister")
            st.divider()
            lbl, cap = _VERTEGEN_INFO.get(route, _VERTEGEN_INFO["bv"])
            st.write(f"**B. {lbl}** (art. 33 WWFT)")
            st.caption(cap + " Géén kopie met BSN/pasfoto bewaren (AVG).")
            persoon_df = st.data_editor(
                _id_df(True), column_config=_id_col_cfg(True),
                num_rows="dynamic", use_container_width=True, key="persoon_editor",
            )
            st.divider()
            _ubo_opties = UBO_GRONDSLAG.get(route, UBO_GRONDSLAG["bv"])
            st.write("**C. UBO-vaststelling**")
            st.caption(_UBO_INFO.get(route, _UBO_INFO["bv"]))
            ubo_df = st.data_editor(
                pd.DataFrame({"Naam": [""], "Grondslag": [_ubo_opties[0]], "Belang / toelichting": [""]}),
                column_config={
                    "Naam": st.column_config.TextColumn("Naam", width="medium"),
                    "Grondslag": st.column_config.SelectboxColumn("Grondslag", options=_ubo_opties, width="large"),
                    "Belang / toelichting": st.column_config.TextColumn("Belang / toelichting", width="medium"),
                },
                num_rows="dynamic", use_container_width=True, key="ubo_editor",
            )
            st.divider()
            st.write("**D. UBO-register controle (kvk.nl/ubo-register)**")
            col1, col2 = st.columns(2)
            ubo_register = col1.checkbox("UBO-register gecontroleerd")
            ubo_datum_input = col2.date_input("Datum controle", value=None, key="ubo_reg_datum")

        submit_personen = st.form_submit_button("Bevestigen en doorgaan →", type="primary")

    if submit_personen:
        vertegenwoordigers = []
        if persoon_df is not None:
            for _, row in persoon_df.iterrows():
                naam = str(row.get("Naam", "")).strip()
                if naam:
                    vertegenwoordigers.append({
                        "naam": naam,
                        "bevoegdheid": str(row.get("Bevoegdheid", "") or ""),
                        "geboortedatum": str(row.get("Geboortedatum", "") or ""),
                        "geboorteplaats": str(row.get("Geboorteplaats", "") or ""),
                        "woonplaats": str(row.get("Woonplaats", "") or ""),
                        "soort_id": str(row.get("Soort ID", "") or ""),
                        "documentnummer": str(row.get("Documentnummer", "") or ""),
                        "datum_uitgifte": str(row.get("Datum uitgifte", "") or ""),
                        "plaats_uitgifte": str(row.get("Plaats uitgifte", "") or ""),
                        "geldig_tot": str(row.get("Geldig tot", "") or ""),
                        "datum_gezien": str(row.get("Datum gezien", "") or ""),
                        "wijze_identificatie": str(row.get("Wijze identificatie", "") or ""),
                    })
        ubos = []
        if ubo_df is not None:
            for _, row in ubo_df.iterrows():
                naam = str(row.get("Naam", "")).strip()
                if naam:
                    ubos.append({
                        "naam": naam,
                        "rol": str(row.get("Rol", "") or ""),
                        "grondslag": str(row.get("Grondslag", "") or ""),
                        "belang": str(row.get("Belang / toelichting", "") or ""),
                    })
        st.session_state["personen_data"] = {
            "route": route,
            "kvk_uittreksel_datum": str(kvk_datum) if kvk_datum else "",
            "bevoegdheid_gecontroleerd": bevoegdheid_check,
            "overeenkomst_geraadpleegd": overeenkomst_check,
            "vertegenwoordigers": vertegenwoordigers,
            "ubos": ubos,
            "ubo_register_gecontroleerd": ubo_register,
            "ubo_register_datum": str(ubo_datum_input) if ubo_datum_input else "",
        }
        st.rerun()
    st.stop()

else:
    pd_data = st.session_state["personen_data"]
    route_saved = pd_data.get("route", "bv")
    vertegens = pd_data.get("vertegenwoordigers", [])
    ubos_saved = pd_data.get("ubos", [])
    _disp_labels = {"vof": "Vennoten", "cv": "Beherend vennoot", "stichting": "Bestuurder(s)", "vereniging": "Bestuurder(s)"}

    if route_saved == "eenmanszaak":
        if vertegens:
            v = vertegens[0]
            st.write(f"**Ondernemer:** {v['naam']}")
            if v.get("soort_id"):
                st.caption(f"{v['soort_id']} {v.get('documentnummer','')} uitgifte {v.get('datum_uitgifte','')} {v.get('plaats_uitgifte','')} — gezien {v.get('datum_gezien','')}")
    else:
        label = _disp_labels.get(route_saved, "Vertegenwoordiger(s)")
        if vertegens:
            st.write(f"**{label}:**")
            for v in vertegens:
                id_info = f" | {v.get('soort_id','')} {v.get('documentnummer','')} gezien {v.get('datum_gezien','')}".rstrip()
                bev = f" ({v['bevoegdheid']})" if v.get("bevoegdheid") else ""
                st.write(f"- **{v['naam']}**{bev}{id_info}")
        if ubos_saved:
            st.write("**UBO's:**")
            for u in ubos_saved:
                rol_info = f" [{u['rol']}]" if u.get("rol") else ""
                st.write(f"- **{u['naam']}**{rol_info} — {u.get('grondslag','')} {u.get('belang','')}")
        ubo_reg = pd_data.get("ubo_register_gecontroleerd")
        ubo_d = pd_data.get("ubo_register_datum", "")
        st.write(f"UBO-register: {'✅ Gecontroleerd' if ubo_reg else '⚠️ Niet gecontroleerd'}{' (' + ubo_d + ')' if ubo_d else ''}")

    if st.button("Wijzigen", key="wijzig_personen"):
        for k in ["personen_data", "screening_bedrijf", "screening_personen", "media", "integriteitsvragen"]:
            st.session_state.pop(k, None)
        st.rerun()

# ──────────────────────────────────────────────
# Stap 5: Screening
# ──────────────────────────────────────────────

st.divider()
st.subheader("🔎 Stap 5: Screening")
st.caption("OpenSanctions — EU, VN, OFAC sanctielijsten + PEP-registers (~40 bronnen)")

if "screening_bedrijf" not in st.session_state:
    with st.spinner(f"Screening bedrijf '{bedrijfsnaam}'…"):
        st.session_state["screening_bedrijf"] = screen_opensanctions(bedrijfsnaam, "Company")

if "screening_personen" not in st.session_state:
    sp = {}
    pd_data = st.session_state.get("personen_data", {})
    te_screenen = {}
    for v in pd_data.get("vertegenwoordigers", []):
        if v.get("naam", "").strip():
            te_screenen[v["naam"]] = True
    for u in pd_data.get("ubos", []):
        naam = u.get("naam", "").strip()
        if naam and "pseudo" not in u.get("grondslag", "").lower():
            te_screenen[naam] = True
    for naam in te_screenen:
        with st.spinner(f"Screening persoon '{naam}'…"):
            sp[naam] = screen_opensanctions(naam, "Person")
    st.session_state["screening_personen"] = sp

screening_bedrijf = st.session_state["screening_bedrijf"]
screening_personen = st.session_state["screening_personen"]

st.write("**Bedrijf:**")
if not screening_bedrijf.get("success"):
    if screening_bedrijf.get("geen_key"):
        st.warning("OpenSanctions API key niet ingesteld — screening overgeslagen. Risicobeoordeling loopt door.")
    else:
        st.error(f"Screening mislukt: {screening_bedrijf.get('error','')}")
else:
    hoog = screening_bedrijf.get("hoog_risico", 0)
    hits = screening_bedrijf.get("hits", 0)
    if hoog > 0:
        st.error(f"⚠️ {hoog} hoge-score treffer(s) — bedrijf op sanctielijst!")
    elif hits > 0:
        st.warning(f"⚠️ {hits} mogelijke treffer(s) — nader onderzoek vereist")
    else:
        st.success("✅ Geen treffers (bedrijf)")
    if screening_bedrijf.get("resultaten"):
        with st.expander("Screeningresultaten bedrijf"):
            for r in screening_bedrijf.get("resultaten", []):
                sv = r.get("score", 0)
                props = r.get("entity", {}).get("properties", {})
                icon = "🔴" if sv >= 0.85 else "🟡"
                naam_r = props.get("name", ["Onbekend"])[0]
                topics = ", ".join(props.get("topics", [])) or "–"
                st.write(f"{icon} **{naam_r}** | Score: {sv:.0%} | {topics}")

if screening_personen:
    st.write("**Personen:**")
    for naam, res in screening_personen.items():
        if not res.get("success"):
            st.write(f"- {naam}: niet gescreend")
        else:
            hoog = res.get("hoog_risico", 0)
            hits = res.get("hits", 0)
            if hoog > 0:
                st.error(f"⚠️ {naam}: {hoog} hoge-score treffer(s)!")
            elif hits > 0:
                st.warning(f"⚠️ {naam}: {hits} mogelijke treffer(s)")
            else:
                st.success(f"✅ {naam}: geen treffers")
            if res.get("resultaten"):
                with st.expander(f"Details {naam}"):
                    for r in res.get("resultaten", []):
                        sv = r.get("score", 0)
                        props = r.get("entity", {}).get("properties", {})
                        icon = "🔴" if sv >= 0.85 else "🟡"
                        naam_r = props.get("name", ["Onbekend"])[0]
                        topics = ", ".join(props.get("topics", [])) or "–"
                        st.write(f"{icon} **{naam_r}** | Score: {sv:.0%} | {topics}")

st.divider()
st.subheader("📰 Adverse media search")
st.caption("DuckDuckGo — fraude, oplichting, witwassen, faillissement, sanctie")

if "media" not in st.session_state:
    with st.spinner(f"Negatief nieuws zoeken over '{bedrijfsnaam}'…"):
        st.session_state["media"] = zoek_adverse_media(bedrijfsnaam)

media = st.session_state["media"]
if not media["success"]:
    st.warning(f"Media search mislukt: {media.get('error','')}")
else:
    media_resultaten = media.get("resultaten", [])
    if not media_resultaten:
        st.success("✅ Geen negatief nieuws gevonden")
    else:
        st.warning(f"⚠️ {len(media_resultaten)} relevante resultaten — beoordeel handmatig")
        for r in media_resultaten:
            with st.expander(r["titel"][:100]):
                st.write(r["tekst"])
                st.markdown(f"[Bekijk artikel]({r['url']})")

# ──────────────────────────────────────────────
# Stap 6: Integriteitsvragen
# ──────────────────────────────────────────────

st.divider()
st.subheader("🔎 Stap 6: Integriteitsvragen")
st.caption("Conform NBA Nadere Voorschriften Witwassen (NV WWFT) — vragen 4 t/m 11")

if "integriteitsvragen" not in st.session_state:
    with st.form("integriteit_form"):
        antwoorden = {}
        for vraag in INTEGRITEITSVRAGEN:
            st.write(f"**{vraag['nr']}. {vraag['vraag']}**")
            if vraag["type"] == "vrij_tekst":
                tekst = st.text_area(
                    "Antwoord", key=f"int_{vraag['nr']}", height=68,
                    label_visibility="collapsed",
                )
                antwoorden[str(vraag["nr"])] = {"tekst": tekst}
            else:
                col1, col2 = st.columns([1, 2])
                waarde = col1.radio(
                    "Antwoord", ["Nee", "Ja", "N.v.t."],
                    key=f"int_{vraag['nr']}", horizontal=True,
                    label_visibility="collapsed",
                )
                toelichting_v = ""
                if waarde == "Ja":
                    toelichting_v = col2.text_input(
                        "Toelichting (verplicht bij Ja)",
                        key=f"int_toe_{vraag['nr']}",
                    )
                    if vraag.get("kritiek"):
                        st.error("⚠️ Direct melden aan compliance officer!")
                antwoorden[str(vraag["nr"])] = {"waarde": waarde, "toelichting": toelichting_v}
            st.markdown("---")

        submit_integriteit = st.form_submit_button("Bevestigen en doorgaan →", type="primary")

    if submit_integriteit:
        st.session_state["integriteitsvragen"] = antwoorden
        st.rerun()
    st.stop()
else:
    antwoorden = st.session_state["integriteitsvragen"]
    ja_nrs = [
        v["nr"] for v in INTEGRITEITSVRAGEN
        if v["type"] == "ja_nee" and antwoorden.get(str(v["nr"]), {}).get("waarde") == "Ja"
    ]
    if ja_nrs:
        st.warning(f"⚠️ Integriteitsrisico's gesignaleerd bij vraag {', '.join(map(str, ja_nrs))}")
    else:
        st.success("✅ Geen integriteitsrisico's gesignaleerd")
    if st.button("Wijzigen", key="wijzig_integriteit"):
        st.session_state.pop("integriteitsvragen", None)
        st.rerun()

# ──────────────────────────────────────────────
# Stap 7: Risicobeoordeling
# ──────────────────────────────────────────────

st.divider()
st.subheader("⚖️ Stap 7: Risicobeoordeling")

land = st.session_state.get("doel_aard", {}).get("land", "Nederland")
risico_klasse, score, factoren = bereken_risico(
    basisprofiel,
    screening_bedrijf,
    screening_personen,
    land,
    st.session_state.get("integriteitsvragen", {}),
)
cdd_vorm, cdd_artikel = get_cdd_vorm(risico_klasse)

col1, col2 = st.columns([1, 2])
with col1:
    st.metric("Risicoclassificatie", f"{RISICO_KLEUR[risico_klasse]} {risico_klasse}")
    st.metric("Vereiste CDD-vorm", cdd_vorm)
    st.caption(cdd_artikel)
    st.metric("Totaalscore", score)
with col2:
    st.write("**Risicobepalende factoren:**")
    if factoren:
        for f in factoren:
            st.write(f"- {f}")
    else:
        st.write("- Geen bijzondere risicofactoren")

if risico_klasse == "HOOG":
    st.error("⚠️ Verscherpt cliëntenonderzoek vereist (art. 8 WWFT). Melding compliance officer verplicht.")
elif risico_klasse == "MIDDEN":
    st.warning("Verhoogde alertheid gedurende de relatie aanbevolen.")

toelichting = st.text_area(
    "Toelichting / motivatie medewerker (optioneel)",
    placeholder="Aanvullende informatie of afwijkende beoordeling…",
    height=80,
)

# ──────────────────────────────────────────────
# Stap 8: Rapport downloaden
# ──────────────────────────────────────────────

st.divider()
st.subheader("📄 Stap 8: Rapport downloaden")
st.caption("Het Word-rapport bevat alle stappen, identificatiegegevens, screeningsuitkomsten en de risicoclassificatie.")

nu = datetime.now()
referentie = f"WWFT-{st.session_state.get('kvk_nummer','?')}-{nu.strftime('%Y%m%d')}"
medewerker = st.session_state.get("doel_aard", {}).get("medewerker", "")
doel_aard = st.session_state.get("doel_aard", {})
personen_data = st.session_state.get("personen_data", [])
integriteit_antwoorden = st.session_state.get("integriteitsvragen", {})

word_bytes = genereer_word(
    kvk_data, basisprofiel, personen_data,
    screening_bedrijf, screening_personen, media,
    integriteit_antwoorden, risico_klasse, score, factoren,
    cdd_vorm, cdd_artikel, doel_aard, toelichting,
    medewerker, referentie, nu,
)

bestandsnaam = f"wwft_{st.session_state.get('kvk_nummer','check')}_{nu.strftime('%Y%m%d_%H%M')}.docx"
st.download_button(
    label="⬇️ Download WWFT-rapport (Word)",
    data=word_bytes,
    file_name=bestandsnaam,
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    type="primary",
)
